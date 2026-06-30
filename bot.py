import os
import sys
import logging
import io
import tempfile
from datetime import datetime
from telegram import Update, InlineKeyboardButton, InlineKeyboardMarkup
from telegram.ext import Application, CommandHandler, MessageHandler, filters, CallbackQueryHandler, ContextTypes

# Document conversion libraries
from docx import Document
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import inch
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
import PyPDF2

# Enable logging
logging.basicConfig(
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    level=logging.INFO
)
logger = logging.getLogger(__name__)

# Get bot token from environment variable
def get_token():
    """Get bot token from environment variables."""
    token = os.environ.get('BOT_TOKEN')
    if not token:
        token = os.environ.get('TELEGRAM_BOT_TOKEN')
    if not token:
        logger.error("❌ No BOT_TOKEN found in environment variables!")
        logger.error("Please add BOT_TOKEN to your Railway Variables.")
        sys.exit(1)
    return token

TOKEN = get_token()
logger.info("✅ Bot token loaded successfully!")

# Store user sessions
user_sessions = {}

# Conversion options
CONVERSION_TYPES = {
    'docx_to_pdf': '📄 Word (.docx) → PDF',
    'pdf_to_docx': '📄 PDF → Word (.docx)',
    'txt_to_pdf': '📝 Text (.txt) → PDF',
    'docx_to_txt': '📄 Word (.docx) → Text (.txt)',
    'pdf_to_txt': '📄 PDF → Text (.txt)'
}

# File size limit (20MB)
MAX_FILE_SIZE = 20 * 1024 * 1024

async def start(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    """Send a welcome message with conversion options."""
    user = update.effective_user
    
    keyboard = [
        [InlineKeyboardButton("📄 Word → PDF", callback_data="convert_docx_to_pdf")],
        [InlineKeyboardButton("📄 PDF → Word", callback_data="convert_pdf_to_docx")],
        [InlineKeyboardButton("📄 Word → Text", callback_data="convert_docx_to_txt")],
        [InlineKeyboardButton("📄 PDF → Text", callback_data="convert_pdf_to_txt")],
        [InlineKeyboardButton("📝 Text → PDF", callback_data="convert_txt_to_pdf")],
        [InlineKeyboardButton("❓ Help", callback_data="help")]
    ]
    reply_markup = InlineKeyboardMarkup(keyboard)
    
    welcome_text = f"""
📁 **Welcome to ConvertMyFile0Bot, {user.first_name}!**

I convert documents between different formats.

**Select a conversion type below:**
"""
    await update.message.reply_text(welcome_text, reply_markup=reply_markup)


async def help_command(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    """Send a help message."""
    help_text = """
📖 **How to use ConvertMyFile0Bot:**

1️⃣ Select a conversion type from the menu
2️⃣ Send me the file to convert
3️⃣ Get your converted file back!

**Supported conversions:**
• Word (.docx) → PDF
• PDF → Word (.docx)
• Word (.docx) → Text (.txt)
• PDF → Text (.txt)
• Text (.txt) → PDF

**Commands:**
/start - Open main menu
/help - Show this help message

💡 **Tip:** Send a file directly and I'll guide you!
"""
    await update.message.reply_text(help_text)


async def button_callback(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    """Handle button presses."""
    query = update.callback_query
    await query.answer()
    
    user_id = update.effective_user.id
    data = query.data
    
    if data == "help":
        await query.edit_message_text(
            "📖 **Help Menu**\n\n"
            "1. Select a conversion type\n"
            "2. Send your file\n"
            "3. Get converted file!\n\n"
            "Use /start to return to main menu."
        )
        return
    
    if data.startswith("convert_"):
        conversion_type = data.replace("convert_", "")
        user_sessions[user_id] = {'conversion': conversion_type}
        
        conversion_name = CONVERSION_TYPES.get(conversion_type, conversion_type)
        
        # Show format-specific instructions
        instructions = {
            'docx_to_pdf': "Please send a **.docx** file to convert to PDF.",
            'pdf_to_docx': "Please send a **.pdf** file to convert to Word (.docx).",
            'docx_to_txt': "Please send a **.docx** file to convert to Text (.txt).",
            'pdf_to_txt': "Please send a **.pdf** file to convert to Text (.txt).",
            'txt_to_pdf': "Please send a **.txt** file to convert to PDF."
        }
        
        await query.edit_message_text(
            f"✅ **Selected: {conversion_name}**\n\n"
            f"{instructions.get(conversion_type, 'Send your file for conversion.')}\n\n"
            f"Send /start to choose a different conversion."
        )


async def convert_docx_to_pdf(docx_data: bytes) -> bytes:
    """Convert .docx to PDF using python-docx and ReportLab."""
    try:
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_docx:
            temp_docx.write(docx_data)
            docx_path = temp_docx.name
        
        doc = Document(docx_path)
        
        pdf_buffer = io.BytesIO()
        c = canvas.Canvas(pdf_buffer, pagesize=A4)
        width, height = A4
        
        y = height - inch
        left_margin = inch
        font_size = 11
        
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                if paragraph.style.name.startswith('Heading'):
                    c.setFont('Helvetica-Bold', font_size + 2)
                else:
                    c.setFont('Helvetica', font_size)
                
                text = paragraph.text
                lines = []
                current_line = ""
                for word in text.split():
                    test_line = current_line + word + " "
                    if len(test_line) * 5 < (width - 2 * left_margin):
                        current_line = test_line
                    else:
                        if current_line:
                            lines.append(current_line.strip())
                        current_line = word + " "
                if current_line:
                    lines.append(current_line.strip())
                
                for line in lines:
                    if y < inch:
                        c.showPage()
                        c.setFont('Helvetica', font_size)
                        y = height - inch
                    c.drawString(left_margin, y, line)
                    y -= (font_size + 2)
                y -= 4
        
        c.save()
        pdf_buffer.seek(0)
        os.unlink(docx_path)
        return pdf_buffer.getvalue()
        
    except Exception as e:
        logger.error(f"DOCX to PDF conversion error: {e}")
        raise e


async def convert_pdf_to_docx(pdf_data: bytes) -> bytes:
    """Convert PDF to .docx by extracting text."""
    try:
        with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as temp_pdf:
            temp_pdf.write(pdf_data)
            pdf_path = temp_pdf.name
        
        doc = Document()
        
        with open(pdf_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            for page in pdf_reader.pages:
                text = page.extract_text()
                if text.strip():
                    doc.add_paragraph(text.strip())
                doc.add_paragraph()
        
        docx_buffer = io.BytesIO()
        doc.save(docx_buffer)
        docx_buffer.seek(0)
        
        os.unlink(pdf_path)
        return docx_buffer.getvalue()
        
    except Exception as e:
        logger.error(f"PDF to DOCX conversion error: {e}")
        raise e


async def convert_docx_to_txt(docx_data: bytes) -> bytes:
    """Convert .docx to .txt by extracting text."""
    try:
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_docx:
            temp_docx.write(docx_data)
            docx_path = temp_docx.name
        
        doc = Document(docx_path)
        
        text_content = ""
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_content += paragraph.text + "\n"
        
        os.unlink(docx_path)
        return text_content.encode('utf-8')
        
    except Exception as e:
        logger.error(f"DOCX to TXT conversion error: {e}")
        raise e


async def convert_pdf_to_txt(pdf_data: bytes) -> bytes:
    """Convert PDF to .txt by extracting text."""
    try:
        with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as temp_pdf:
            temp_pdf.write(pdf_data)
            pdf_path = temp_pdf.name
        
        text_content = ""
        with open(pdf_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            for page in pdf_reader.pages:
                text = page.extract_text()
                if text.strip():
                    text_content += text + "\n\n"
        
        os.unlink(pdf_path)
        return text_content.encode('utf-8')
        
    except Exception as e:
        logger.error(f"PDF to TXT conversion error: {e}")
        raise e


async def convert_txt_to_pdf(txt_data: bytes) -> bytes:
    """Convert .txt to PDF using ReportLab."""
    try:
        text_content = txt_data.decode('utf-8')
        
        pdf_buffer = io.BytesIO()
        c = canvas.Canvas(pdf_buffer, pagesize=A4)
        width, height = A4
        
        y = height - inch
        left_margin = inch
        font_size = 11
        c.setFont('Helvetica', font_size)
        
        lines = text_content.split('\n')
        for line in lines:
            if y < inch:
                c.showPage()
                c.setFont('Helvetica', font_size)
                y = height - inch
            c.drawString(left_margin, y, line.strip() if line.strip() else ' ')
            y -= (font_size + 2)
        
        c.save()
        pdf_buffer.seek(0)
        return pdf_buffer.getvalue()
        
    except Exception as e:
        logger.error(f"TXT to PDF conversion error: {e}")
        raise e


async def handle_document(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    """Handle document files sent to the bot."""
    user_id = update.effective_user.id
    document = update.message.document
    
    # Check file size
    if document.file_size and document.file_size > MAX_FILE_SIZE:
        await update.message.reply_text(
            f"❌ **File too large!**\n\n"
            f"Maximum file size: 20MB\n"
            f"Your file: {document.file_size / 1024 / 1024:.1f}MB\n\n"
            f"Please send a smaller file."
        )
        return
    
    # Check if user has selected a conversion
    if user_id not in user_sessions or 'conversion' not in user_sessions[user_id]:
        await update.message.reply_text(
            "❌ Please select a conversion type first!\n\n"
            "Use /start to open the main menu and choose a conversion."
        )
        return
    
    conversion_type = user_sessions[user_id]['conversion']
    file_name = document.file_name or "document"
    
    # Validate file type
    valid_extensions = {
        'docx_to_pdf': '.docx',
        'pdf_to_docx': '.pdf',
        'docx_to_txt': '.docx',
        'pdf_to_txt': '.pdf',
        'txt_to_pdf': '.txt'
    }
    
    expected_ext = valid_extensions.get(conversion_type)
    if expected_ext and not file_name.lower().endswith(expected_ext):
        await update.message.reply_text(
            f"❌ **Wrong file type!**\n\n"
            f"I need a {expected_ext} file for this conversion.\n"
            f"Received: `{file_name}`\n\n"
            f"Please send the correct file type."
        )
        return
    
    # Send processing message
    processing_msg = await update.message.reply_text(
        f"🔄 **Converting your file...**\n\n"
        f"📄 File: `{file_name}`\n"
        f"⏳ This may take a moment..."
    )
    
    try:
        # Download the file
        file = await document.get_file()
        file_data = await file.download_as_bytearray()
        
        # Perform conversion
        if conversion_type == 'docx_to_pdf':
            converted_data = await convert_docx_to_pdf(file_data)
            output_ext = '.pdf'
            conversion_name = 'PDF'
        elif conversion_type == 'pdf_to_docx':
            converted_data = await convert_pdf_to_docx(file_data)
            output_ext = '.docx'
            conversion_name = 'Word (.docx)'
        elif conversion_type == 'docx_to_txt':
            converted_data = await convert_docx_to_txt(file_data)
            output_ext = '.txt'
            conversion_name = 'Text (.txt)'
        elif conversion_type == 'pdf_to_txt':
            converted_data = await convert_pdf_to_txt(file_data)
            output_ext = '.txt'
            conversion_name = 'Text (.txt)'
        elif conversion_type == 'txt_to_pdf':
            converted_data = await convert_txt_to_pdf(file_data)
            output_ext = '.pdf'
            conversion_name = 'PDF'
        else:
            raise ValueError(f"Unknown conversion type: {conversion_type}")
        
        # Generate output filename
        base_name = file_name.rsplit('.', 1)[0]
        output_filename = f"{base_name}_converted{output_ext}"
        
        # Send the converted file back
        await update.message.reply_document(
            document=io.BytesIO(converted_data),
            filename=output_filename,
            caption=f"✅ **Converted successfully!**\n\n"
                    f"📄 Original: `{file_name}`\n"
                    f"📄 Converted: `{output_filename}`\n"
                    f"📊 Size: {len(converted_data) / 1024:.1f} KB\n"
                    f"🔄 Format: {conversion_name}"
        )
        
        await processing_msg.delete()
        
        # Reset user session
        del user_sessions[user_id]
        
    except Exception as e:
        logger.error(f"Conversion failed: {e}")
        await processing_msg.edit_text(
            f"❌ **Sorry, I couldn't convert that file.**\n\n"
            f"Error: {str(e)}\n\n"
            f"💡 **Tips:**\n"
            f"• Make sure it's a valid file\n"
            f"• Try a different file\n"
            f"• Use /start to select a different conversion"
        )


async def handle_text(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    """Handle text messages that aren't commands."""
    await update.message.reply_text(
        "📤 Please use the menu to select a conversion type first.\n\n"
        "Send /start to open the main menu."
    )


def main() -> None:
    """Start the bot."""
    try:
        # Create Application
        application = Application.builder().token(TOKEN).build()
        
        # Add command handlers
        application.add_handler(CommandHandler("start", start))
        application.add_handler(CommandHandler("help", help_command))
        
        # Add callback handler for inline buttons
        application.add_handler(CallbackQueryHandler(button_callback))
        
        # Add message handlers
        application.add_handler(MessageHandler(filters.Document.ALL, handle_document))
        application.add_handler(MessageHandler(filters.TEXT & ~filters.COMMAND, handle_text))
        
        # Start the Bot
        logger.info("🚀 ConvertMyFile0Bot started successfully!")
        logger.info("📁 Press Ctrl+C to stop.")
        application.run_polling()
        
    except Exception as e:
        logger.error(f"❌ Fatal error: {e}")
        sys.exit(1)


if __name__ == '__main__':
    main()
